# imports
from tkinter import *
import tkinter as tk
import pandas as pd
import xlwings as xw
# from displayenteredinfo import displayenteredinfo
# from location import *
# gui
root = tk.Tk()
# tkinter functions go here
def close_window():
    root.destroy()    
def submit(): 
    srs=list(e1.get().split(','))
    contingency=list(e2.get().split(','))
    gst=list(e3.get().split(','))
    location=var.get()
    print (srs,contingency,gst,location)
    displayenteredinfo(srs,contingency,gst,location)
    dataframeList=excel(srs,contingency,gst,location)
    wordProc(dataframeList)
    return srs,contingency,gst,location
# tkinter widgets code goes here
l1=Label(root,text="Srs number")
l1.place(x=10,y=10)
e1=Entry(root,bd=5,width=100)
e1.place(x=80,y=10)
l2=Label(root,text="contingency")
l2.place(x=10,y=30)
e2=Entry(root,bd=5)
e2.place(x=80,y=30)
l3=Label(root,text="gst")
l3.place(x=10,y=50)
e3=Entry(root,bd=5)
e3.place(x=80,y=50)
var = StringVar(root)
option = OptionMenu(root, var, "l", "rural", "urban")
option.place(x=10,y=80)
close=tk.Button(root,text='Close',command=close_window)
close.place(x=150,y=120)
submit=tk.Button(root,text='Submit',command=submit)
submit.place(x=100,y=120)
# gui

def excel(srs,contingency,gst,loc):
    excel_file='PWD Electrical CSR18-19 Final.xlsx'
    sheetname='Sheet1'
    wb = xw.Book(excel_file)
    pwd = pd.read_excel(excel_file)
    sht = wb.sheets[sheetname]
    dataframeList=[]
    for i in range(len(srs)):
        myCell = wb.sheets[sheetname].api.UsedRange.Find(srs[i]) #srs number will go here
        Address = str(myCell.address)[3:]
        # main list of alues of columns hence extracted
        dataframe = pwd.iloc[int(Address)-2:int(Address)-1].values.tolist()
        dataframe = list(dataframe[0])
        dataframe = dataframe[:6]
        dataframeList.append(dataframe)
    return dataframeList
def displayenteredinfo(srs,contingency,gst,loc):
    root = Tk()
    label = Label( root, text=srs)
    label.place(x=10,y=10)
    label1 = Label( root, text=contingency)
    label1.place(x=10,y=30)
    label2 = Label( root, text=gst)
    label2.place(x=10,y=50)
    label3 = Label( root, text=loc)
    label3.place(x=10,y=70)
    root.mainloop()
def wordProc(dataframeList):
    from docx import Document
    from docx.shared import Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    document = Document()
    document.add_heading('ABSTRACT',0)
    # main body
    document.add_paragraph('Estimate no: ')
    document.add_paragraph('Name of work: ')
    # tables
    # sample records actual records to be taken from  sheet itself
    records=(dataframeList)
    table=document.add_table(rows=1,cols=6)
    hdr_cells=table.rows[0].cells
    hdr_cells[0].text='Srno'
    hdr_cells[1].text='qty'
    hdr_cells[2].text='desc of item'
    hdr_cells[3].text='contingency'
    hdr_cells[4].text='gst'
    hdr_cells[5].text='amount'
    for Sr_no,qty,desc_of_item,contingency,gst,amount in records:

        row_cells=table.add_row().cells

        row_cells[0].text=str(Sr_no)
        row_cells[1].text=str(qty)
        row_cells[2].text=str(desc_of_item)
        row_cells[3].text=str(contingency)
        row_cells[4].text=str(gst)
        row_cells[5].text=str(amount)
            


    # final calculations
    document.add_paragraph(' ')
    p1=document.add_paragraph('total rs: ')
    p1.paragraph_format.alignment=WD_ALIGN_PARAGRAPH.RIGHT
    p2=document.add_paragraph('add 4% for contigency rs: ')
    p2.paragraph_format.alignment=WD_ALIGN_PARAGRAPH.RIGHT
    p3=document.add_paragraph('total rs: ')
    p3.paragraph_format.alignment=WD_ALIGN_PARAGRAPH.RIGHT
    p4=document.add_paragraph('add 12% for gst rs: ')
    p4.paragraph_format.alignment=WD_ALIGN_PARAGRAPH.RIGHT
    p5=document.add_paragraph('total rs: ')
    p5.paragraph_format.alignment=WD_ALIGN_PARAGRAPH.RIGHT
    p6=document.add_paragraph('say rs: ')
    p6.paragraph_format.alignment=WD_ALIGN_PARAGRAPH.RIGHT

    document.add_paragraph(' ')


    # signatures
    document.add_paragraph('100% arithmetically checked by me')
    document.add_paragraph(' ')
    document.add_paragraph(' ')

    # footer
    document.add_page_break()
    report=('report.docx')
    document.save(report)
    # header
        

root.mainloop()